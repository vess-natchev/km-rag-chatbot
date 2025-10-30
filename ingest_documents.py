#!/usr/bin/env python3
"""
Comprehensive Document Ingestion Script for Org Knowledge Management System

This script processes Word (.docx), Excel (.xlsx), and PDF files,
extracting text content and rich metadata for the knowledge base.

Key Features:
- Multi-format support (DOCX, XLSX, PDF)
- Multi-language support (English, Spanish, French)
- Automatic metadata extraction (country, status, project details)
- Status detection (approved/rejected) from folder structure
- Semantic chunking based on document structure
- Robust error handling
- Batch processing with progress tracking
"""

import asyncio
import os
import sys
import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, cast
import aiohttp
import argparse
from datetime import datetime
import os
from config import settings
import nltk

# Make .env variables available to any downstream libraries (like Unstructured)
os.environ.update({k: str(v) for k, v in settings.model_dump().items() if v is not None})

if os.getenv("AUTO_DOWNLOAD_NLTK", "").lower() in ("1", "true", "yes"):
    # Correct modern resource names
    for pkg, path in [
        ("punkt_tab", "tokenizers/punkt_tab"),
        ("averaged_perceptron_tagger_eng", "taggers/averaged_perceptron_tagger_eng"),
    ]:
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(pkg, quiet=True)

nltk.download('averaged_perceptron_tagger')

# Document parsing libraries
from docx import Document as DocxDocument
from openpyxl import load_workbook
import pymupdf  # PyMuPDF for PDF processing

# Unstructured library for better document parsing
from unstructured.partition.auto import partition
from unstructured.partition.docx import partition_docx
from unstructured.partition.doc import partition_doc
from unstructured.partition.xlsx import partition_xlsx
from unstructured.partition.pdf import partition_pdf

# Language configuration for multi-language document processing
# ISO 639-3 language codes for Unstructured library
DOCUMENT_LANGUAGES = ["eng", "spa", "fra"]  # English, Spanish, French

# Status keywords for detection
APPROVED_KEYWORDS = [
    'approved', 'awarded', 'funded', 'successful', 'accepted',
    'completed', 'implemented', 'active', 'ongoing'
]
REJECTED_KEYWORDS = [
    'rejected', 'declined', 'unsuccessful', 'denied',
    'not approved', 'not selected', 'withdrawn'
]

# Country list for detection
ORG_COUNTRIES = [
    'kenya', 'uganda', 'rwanda', 'tanzania', 'ethiopia', 'malawi',
    'zambia', 'zimbabwe', 'senegal', 'nigeria', 'ghana',
    'guatemala', 'honduras', 'mexico', 'haiti', 'nicaragua', 'ecuador',
    'bangladesh', 'nepal', 'cambodia', 'india', 'philippines'
]

# Value chain keywords
VALUE_CHAINS = [
    'livestock', 'dairy', 'poultry', 'cattle', 'goats', 'sheep', 'pigs',
    'crops', 'vegetables', 'fruits', 'coffee', 'cocoa', 'honey',
    'aquaculture', 'fish farming', 'agroforestry'
]

class DocumentProcessor:
    """Processes various document formats and extracts metadata with structure preservation."""
    
    def __init__(self, file_path: Path):
        self.file_path = file_path
        self.file_name = file_path.name
        self.file_ext = file_path.suffix.lower()
        
    def extract_text_and_metadata(self) -> Optional[Dict[str, Any]]:
        """
        Extract text and metadata based on file type.
        Uses Unstructured library to preserve document structure.
        """
        try:
            if self.file_ext == '.docx':
                return self._process_docx_structured()
            elif self.file_ext == '.doc':
                return self._process_doc_structured()
            elif self.file_ext in ('.xlsx', '.xls'):
                return self._process_xlsx_structured()
            elif self.file_ext == '.pdf':
                return self._process_pdf_structured()
            else:
                return self._process_generic()
        except Exception as e:
            print(f"Error processing {self.file_path}: {e}")
            # Fallback to basic extraction
            return self._fallback_extract()
    
    def _process_docx_structured(self) -> Optional[Dict[str, Any]]:
        """
        Process Word document with structure preservation.
        Uses Unstructured library to maintain headings and paragraphs.
        """
        try:
            # Use Unstructured library for better structure extraction
            # IMPORTANT: Specify languages to avoid warning and ensure proper text extraction
            elements = partition_docx(
                str(self.file_path),
                languages=DOCUMENT_LANGUAGES  # Multi-language support
            )
            
            # Reconstruct text with preserved structure
            text_parts = []
            for element in elements:
                element_type = type(element).__name__
                text = str(element)
                
                # Preserve structure with newlines
                if element_type == 'Title':
                    text_parts.append(f"\n\n{text}\n")
                elif element_type in ['NarrativeText', 'Text']:
                    text_parts.append(text)
                elif element_type == 'ListItem':
                    text_parts.append(f"â€¢ {text}")
                else:
                    text_parts.append(text)
            
            text_content = '\n'.join(text_parts)
            
            # Also get metadata from docx properties
            doc = DocxDocument(str(self.file_path))
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or self._extract_title_from_filename(),
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
            }
            
            return {
                'content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            print(f"  Structured processing failed, using fallback: {e}")
            return self._fallback_extract_docx()
    
    def _process_doc_structured(self) -> Optional[Dict[str, Any]]:
        """Process legacy Word (.doc) with structure preservation."""
        try:
            # Specify languages for proper text extraction
            elements = partition_doc(
                str(self.file_path),
                languages=DOCUMENT_LANGUAGES  # Multi-language support
            )
            text_content = '\n'.join(str(el) for el in elements)
            title = next((str(el) for el in elements if type(el).__name__ == 'Title'), None)
            metadata = {'title': title or self._extract_title_from_filename()}
            return {'content': text_content, 'metadata': metadata}
        except Exception as e:
            print(f"  Structured .doc processing failed, using generic fallback: {e}")
            return self._fallback_extract()

    def _process_xlsx_structured(self) -> Optional[Dict[str, Any]]:
        """Process Excel spreadsheet with structure preservation."""
        try:
            # Use Unstructured library with language specification
            elements = partition_xlsx(
                str(self.file_path),
                languages=DOCUMENT_LANGUAGES  # Multi-language support
            )
            
            # Reconstruct text with preserved structure
            text_parts = []
            current_sheet = None
            
            for element in elements:
                element_type = type(element).__name__
                text = str(element)
                
                # Track sheet changes
                metadata = element.metadata if hasattr(element, 'metadata') else None
                if metadata and hasattr(metadata, 'page_name'):
                    sheet_name = metadata.page_name
                    if sheet_name != current_sheet:
                        current_sheet = sheet_name
                        text_parts.append(f"\n\n=== Sheet: {sheet_name} ===\n")
                
                # Add element text
                if element_type == 'Title':
                    text_parts.append(f"\n{text}\n")
                else:
                    text_parts.append(text)
            
            text_content = '\n'.join(text_parts)
            
            # Get metadata
            wb = load_workbook(str(self.file_path), read_only=True, data_only=True)
            props = wb.properties
            metadata = {
                'title': props.title or self._extract_title_from_filename(),
                'author': props.creator,
                'subject': props.subject,
                'keywords': props.keywords,
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
            }
            wb.close()
            
            return {
                'content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            print(f"  Structured processing failed, using fallback: {e}")
            return self._fallback_extract_xlsx()
    
    def _process_pdf_structured(self) -> Optional[Dict[str, Any]]:
        """Process PDF document with structure preservation."""
        try:
            # Use Unstructured library for better structure extraction
            # CRITICAL: Specify languages for proper OCR and text extraction
            elements = partition_pdf(
                str(self.file_path),
                strategy="fast",  # Use fast strategy for speed; change to "hi_res" for better quality
                languages=DOCUMENT_LANGUAGES,  # Multi-language support for text extraction
                # For OCR-heavy documents, you can also specify ocr_languages
                # ocr_languages="eng+spa+fra"  # Tesseract format for OCR
            )
            
            # Reconstruct text with preserved structure
            text_parts = []
            for element in elements:
                element_type = type(element).__name__
                text = str(element)
                
                # Preserve structure with newlines
                if element_type == 'Title':
                    text_parts.append(f"\n\n{text}\n")
                elif element_type in ['NarrativeText', 'Text']:
                    text_parts.append(text)
                elif element_type == 'ListItem':
                    text_parts.append(f"â€¢ {text}")
                else:
                    text_parts.append(text)
            
            text_content = '\n'.join(text_parts)
            
            # Get PDF metadata
            doc = pymupdf.open(str(self.file_path))
            pdf_metadata: Dict[str, Any] = cast(Dict[str, Any], doc.metadata or {})
            metadata = {
                'title': pdf_metadata.get('title') or self._extract_title_from_filename(),
                'author': pdf_metadata.get('author'),
                'subject': pdf_metadata.get('subject'),
                'keywords': pdf_metadata.get('keywords'),
                'created': pdf_metadata.get('creationDate'),
                'modified': pdf_metadata.get('modDate'),
            }
            doc.close()
            
            return {
                'content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            print(f"  Structured processing failed, using fallback: {e}")
            return self._fallback_extract_pdf()
    
    def _fallback_extract_docx(self) -> Optional[Dict[str, Any]]:
        """Fallback method for DOCX extraction without structure."""
        doc = DocxDocument(str(self.file_path))
        text_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title or self._extract_title_from_filename(),
            'author': core_props.author,
            'created': core_props.created.isoformat() if core_props.created else None,
        }
        
        return {'content': text_content, 'metadata': metadata}
    
    def _fallback_extract_xlsx(self) -> Optional[Dict[str, Any]]:
        """Fallback method for XLSX extraction without structure."""
        wb = load_workbook(str(self.file_path), read_only=True, data_only=True)
        
        text_parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}\n")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        
        text_content = '\n'.join(text_parts)
        
        props = wb.properties
        metadata = {
            'title': props.title or self._extract_title_from_filename(),
            'author': props.creator,
        }
        wb.close()
        
        return {'content': text_content, 'metadata': metadata}
    
    def _fallback_extract_pdf(self) -> Optional[Dict[str, Any]]:
        """Fallback method for PDF extraction without structure."""
        doc = pymupdf.open(str(self.file_path))
        
        text_parts = []
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        text_content = '\n'.join(text_parts)
        
        pdf_metadata: Dict[str, Any] = cast(Dict[str, Any], doc.metadata or {})
        metadata = {
            'title': pdf_metadata.get('title') or self._extract_title_from_filename(),
            'author': pdf_metadata.get('author'),
        }
        doc.close()
        
        return {'content': text_content, 'metadata': metadata}
    
    def _fallback_extract(self) -> Optional[Dict[str, Any]]:
        """Generic fallback extraction."""
        try:
            # Use generic partition with language specification
            elements = partition(
                str(self.file_path),
                languages=DOCUMENT_LANGUAGES  # Multi-language support
            )
            text_content = '\n'.join([str(el) for el in elements])
            
            return {
                'content': text_content,
                'metadata': {'title': self._extract_title_from_filename()}
            }
        except Exception as e:
            print(f"  All extraction methods failed: {e}")
            return None
    
    def _process_generic(self) -> Optional[Dict[str, Any]]:
        """Process generic files using unstructured library."""
        return self._fallback_extract()
    
    def _extract_title_from_filename(self) -> str:
        """Extract a readable title from filename."""
        title = self.file_path.stem
        title = re.sub(r'[_-]', ' ', title)
        return title

class MetadataExtractor:
    """Extracts structured metadata from file paths and content."""
    
    @staticmethod
    def extract_from_path(file_path: Path) -> Dict[str, Any]:
        """Extract metadata from file path structure."""
        parts = file_path.parts
        metadata: Dict[str, Any] = {}
        
        # Detect status from path (approved vs rejected folders)
        for part in parts:
            part_lower = part.lower()
            if part_lower == 'approved':
                metadata['status'] = 'approved'
                break
            elif part_lower == 'rejected':
                metadata['status'] = 'rejected'
                break
        
        # If not found in simple folder names, check for keywords
        if 'status' not in metadata:
            for part in parts:
                part_lower = part.lower()
                if any(keyword in part_lower for keyword in APPROVED_KEYWORDS):
                    metadata['status'] = 'approved'
                    break
                elif any(keyword in part_lower for keyword in REJECTED_KEYWORDS):
                    metadata['status'] = 'rejected'
                    break
        
        # Detect country from path
        for part in parts:
            part_lower = part.lower()
            for country in ORG_COUNTRIES:
                if country in part_lower:
                    metadata['country'] = part.title()
                    break
        
        # Extract year from filename or path
        for part in parts:
            year_match = re.search(r'\b(20\d{2})\b', part)
            if year_match:
                metadata['year'] = int(year_match.group(1))
                break
        
        # Extract project ID from filename
        filename = file_path.stem
        project_id_patterns = [
            r'(?:PRJ|PROJECT|PROJ)[-_]?(\d+)',
            r'P[-_](\d+)',
            r'(?:ID|NO)[-_]?(\d+)'
        ]
        for pattern in project_id_patterns:
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                metadata['project_id'] = match.group(0)
                break
        
        # Detect document type from filename
        doc_types = {
            'report': ['report', 'rpt'],
            'proposal': ['proposal', 'prop'],
            'evaluation': ['evaluation', 'eval'],
            'budget': ['budget', 'financial'],
            'baseline': ['baseline'],
            'endline': ['endline'],
            'agreement': ['agreement', 'contract', 'mou']
        }
        
        filename_lower = filename.lower()
        for doc_type, keywords in doc_types.items():
            if any(keyword in filename_lower for keyword in keywords):
                metadata['document_type'] = doc_type
                break
        
        return metadata
    
    @staticmethod
    def extract_from_content(content: str, existing_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract additional metadata from content."""
        metadata = existing_metadata.copy()
        
        # Extract value chain from content if not already set
        if 'value_chain' not in metadata:
            content_lower = content.lower()
            for value_chain in VALUE_CHAINS:
                if value_chain in content_lower:
                    metadata['value_chain'] = value_chain
                    break
        
        # Extract project name
        if 'project_name' not in metadata:
            project_name_patterns = [
                r'Project\s+Name:?\s+([^\n]+)',
                r'Project:?\s+([^\n]+)',
                r'Programme\s+Name:?\s+([^\n]+)'
            ]
            for pattern in project_name_patterns:
                match = re.search(pattern, content, re.IGNORECASE)
                if match:
                    metadata['project_name'] = match.group(1).strip()
                    break
        
        # Refine status based on content
        if 'status' not in metadata or metadata['status'] is None:
            content_lower = content[:2000].lower()
            if any(keyword in content_lower for keyword in APPROVED_KEYWORDS):
                metadata['status'] = 'approved'
            elif any(keyword in content_lower for keyword in REJECTED_KEYWORDS):
                metadata['status'] = 'rejected'
        
        return metadata

async def ingest_documents_from_folder(
    folder_path: str,
    api_base_url: str = "http://localhost:8001",
    batch_size: int = 5
):
    """Ingest all documents from a folder."""
    
    folder = Path(folder_path)
    if not folder.exists():
        print(f"Error: Folder {folder_path} does not exist.")
        return
    
    # Find all supported files
    supported_extensions = ['.docx', '.xlsx', '.pdf']
    all_files = []
    for ext in supported_extensions:
        all_files.extend(folder.glob(f"**/*{ext}"))
    
    if not all_files:
        print(f"No supported documents found in {folder_path}")
        print(f"Looking for: {', '.join(supported_extensions)}")
        return
    
    print(f"Found {len(all_files)} documents to ingest...")
    print(f"Supported formats: {', '.join(supported_extensions)}")
    print(f"Multi-language support: English, Spanish, French")
    print(f"Using semantic chunking for structure-aware processing")
    print("="*80)
    
    documents = []
    failed_files = []
    
    for file_path in all_files:
        try:
            print(f"Processing: {file_path.name}...")
            
            # Process document with structure preservation and multi-language support
            processor = DocumentProcessor(file_path)
            doc_data = processor.extract_text_and_metadata()
            
            if not doc_data or not doc_data.get('content'):
                print(f"  âš ï¸  Warning: No content extracted from {file_path.name}")
                failed_files.append((file_path, "No content extracted"))
                continue
            
            # Extract metadata from path
            path_metadata = MetadataExtractor.extract_from_path(file_path)
            
            # Combine metadata
            combined_metadata = {**doc_data.get('metadata', {}), **path_metadata}
            
            # Extract additional metadata from content
            final_metadata = MetadataExtractor.extract_from_content(
                doc_data['content'],
                combined_metadata
            )
            
            # Create document object
            document = {
                "title": final_metadata.get('title', file_path.stem),
                "content": doc_data['content'],
                "file_path": str(file_path.relative_to(folder)),
                "country": final_metadata.get('country'),
                "project_id": final_metadata.get('project_id'),
                "project_name": final_metadata.get('project_name'),
                "project_type": final_metadata.get('project_type'),
                "value_chain": final_metadata.get('value_chain'),
                "document_type": final_metadata.get('document_type'),
                "status": final_metadata.get('status'),
                "year": final_metadata.get('year'),
                "author": final_metadata.get('author'),
                "created_at": final_metadata.get('created'),
                "metadata": final_metadata
            }
            
            documents.append(document)
            
            # Show extracted metadata
            print(f"  âœ" Extracted: {len(doc_data['content'])} chars")
            if document.get('country'):
                print(f"    Country: {document['country']}")
            if document.get('status'):
                print(f"    Status: {document['status']}")
            if document.get('year'):
                print(f"    Year: {document['year']}")
            
        except Exception as e:
            print(f"  âœ— Error: {e}")
            failed_files.append((file_path, str(e)))
            continue
    
    if not documents:
        print("\nâŒ No documents were successfully processed.")
        return
    
    print("\n" + "="*80)
    print(f"Successfully processed {len(documents)} documents")
    if failed_files:
        print(f"Failed to process {len(failed_files)} documents")
    print("="*80)
    
    # Send documents to API in batches
    total_batches = (len(documents) + batch_size - 1) // batch_size
    
    async with aiohttp.ClientSession() as session:
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i + batch_size]
            batch_num = i // batch_size + 1
            
            print(f"\nIngesting batch {batch_num}/{total_batches} ({len(batch)} documents)...")
            
            try:
                async with session.post(
                    f"{api_base_url}/ingest/documents",
                    json={"documents": batch},
                    headers={"Content-Type": "application/json"},
                    timeout=aiohttp.ClientTimeout(total=300)  # 5 min timeout
                ) as response:
                    
                    if response.status == 200:
                        result = await response.json()
                        print(f"  âœ" Batch {batch_num} completed successfully")
                        
                        # Show details
                        for doc_result in result.get('results', []):
                            if doc_result.get('status') == 'success':
                                print(f"    âœ" {doc_result['title']}: {doc_result['embeddings']} embeddings")
                    else:
                        error_text = await response.text()
                        print(f"  âœ— Batch {batch_num} failed: {response.status} - {error_text}")
                        
            except Exception as e:
                print(f"  âœ— Batch {batch_num} failed with error: {e}")
            
            # Small delay between batches
            await asyncio.sleep(2)
    
    print("\n" + "="*80)
    print("âœ… Ingestion process completed!")
    print(f"   Total documents: {len(documents)}")
    print(f"   Successful: {len(documents) - len(failed_files)}")
    print(f"   Failed: {len(failed_files)}")
    print(f"   Languages: English, Spanish, French")
    print(f"   Chunking: Semantic (structure-aware)")
    
    if failed_files:
        print("\nâš ï¸  Failed files:")
        for file_path, error in failed_files:
            print(f"   - {file_path.name}: {error}")

async def main():
    parser = argparse.ArgumentParser(
        description="Ingest Word, Excel, and PDF documents into Org Knowledge Base with multi-language support"
    )
    parser.add_argument("folder", help="Path to folder containing documents")
    parser.add_argument(
        "--api-url",
        default="http://localhost:8001",
        help="Base URL of the RAG API (default: http://localhost:8001)"
    )
    parser.add_argument(
        "--batch-size",
        type=int,
        default=5,
        help="Number of documents per batch (default: 5)"
    )
    
    args = parser.parse_args()
    
    print("="*80)
    print("ORG KNOWLEDGE MANAGEMENT SYSTEM")
    print("Document Ingestion with Multi-Language Support")
    print("="*80)
    print(f"Languages: English, Spanish, French (ISO 639-3: eng, spa, fra)")
    print(f"Strategy: Semantic chunking with structure preservation")
    print("="*80 + "\n")
    
    await ingest_documents_from_folder(
        args.folder,
        args.api_url,
        args.batch_size
    )

if __name__ == "__main__":
    asyncio.run(main())
